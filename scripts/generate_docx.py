#!/usr/bin/env python3
"""Generate .docx report EXACTLY matching EMBUN PAGI EDU DAYCARE format.
Uses exact same logo, fonts, colors, column widths, and layout."""

import sys
import os
import sqlite3
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, Twips, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
ASSETS_DIR = os.path.join(SCRIPT_DIR, 'assets')

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color}" w:themeFill="accent1" w:themeFillTint="66"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_width_twips(cell, twips):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{twips}" w:type="dxa"/>')
    tcPr.append(tcW)

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)
    # Set table style
    tblStyle = parse_xml(f'<w:tblStyle {nsdecls("w")} w:val="TableGrid"/>')
    tblPr.insert(0, tblStyle)

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

def add_para(doc, text, bold=False, size=12, align='left', font='Times New Roman', 
             space_after=0, space_before=0, italic=False):
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    return p

def add_run_to_para(para, text, bold=False, size=12, font='Times New Roman', italic=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    return run

def main():
    if len(sys.argv) < 3:
        print("Usage: generate_docx.py <student_id> <month> [output_path]")
        sys.exit(1)

    student_id = sys.argv[1]
    month = sys.argv[2]
    output_path = sys.argv[3] if len(sys.argv) > 3 else f"/tmp/laporan_{student_id}_{month}.docx"

    db_path = '/root/les-private/les.db'
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row

    student = conn.execute("SELECT * FROM students WHERE id=?", (student_id,)).fetchone()
    if not student:
        print("Student not found")
        sys.exit(1)

    meetings = conn.execute("""
        SELECT * FROM meetings 
        WHERE student_id=? AND strftime('%Y-%m', date)=? AND status='selesai'
        ORDER BY date
    """, (student_id, month)).fetchall()

    # Load assessments grouped by meeting
    meeting_ids = [m['id'] for m in meetings]
    assessments = {}
    if meeting_ids:
        placeholders = ','.join('?' * len(meeting_ids))
        rows = conn.execute(f"""
            SELECT meeting_id, aspect, score, COALESCE(kegiatan,'') as kegiatan 
            FROM assessments WHERE meeting_id IN ({placeholders})
        """, meeting_ids).fetchall()
        for row in rows:
            mid = row['meeting_id']
            asp = row['aspect']
            if mid not in assessments:
                assessments[mid] = {}
            if asp not in assessments[mid]:
                assessments[mid][asp] = []
            assessments[mid][asp].append({
                'kegiatan': row['kegiatan'] if row['kegiatan'] else '-',
                'score': row['score']
            })

    # Summary
    aspect_keys = ['pra_membaca', 'menulis', 'berhitung', 'sensory_play', 'kreativitas', 'brain_game']
    aspect_labels = ['Pra membaca', 'Menulis', 'Berhitung', 'Sensory play', 'Kreativitas', 'Brain game']
    summary = {k: {'BB': 0, 'MB': 0, 'BSH': 0, 'BSB': 0, 'Total': 0, 'kegiatan_list': []} for k in aspect_keys}
    
    for mid in meeting_ids:
        if mid in assessments:
            for k in aspect_keys:
                if k in assessments[mid]:
                    for entry in assessments[mid][k]:
                        score = entry['score']
                        keg = entry['kegiatan']
                        if keg and keg != '-':
                            summary[k]['kegiatan_list'].append(keg.strip())
                        if score == 'Belum Berkembang':
                            summary[k]['BB'] += 1
                        elif score == 'Mulai Berkembang':
                            summary[k]['MB'] += 1
                        elif score == 'Berkembang Sesuai Harapan':
                            summary[k]['BSH'] += 1
                        elif score == 'Berkembang Sangat Baik':
                            summary[k]['BSB'] += 1
                        if score:
                            summary[k]['Total'] += 1

    conn.close()

    # === CREATE DOCUMENT ===
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    for section in doc.sections:
        section.top_margin = Twips(1440)
        section.bottom_margin = Twips(1440)
        section.left_margin = Twips(1440)
        section.right_margin = Twips(1440)

    # === TITLE (langsung tanpa header institusi) ===
    add_para(doc, 'LAPORAN PENILAIAN PERKEMBANGAN', bold=True, size=12, align='center', space_after=6)

    # === STUDENT INFO ===
    add_para(doc, f'Nama: {student["name"]}', size=12, space_after=0)
    add_para(doc, f'Program: Private Tutoring', size=12, space_after=6)

    # === PER MEETING ===
    # Exact column widths from original: [1548, 5616, 818, 595, 754, 633] twips
    col_widths = [1548, 5616, 818, 595, 754, 633]
    header_color = 'B4C6E7'  # Light blue from original

    for m in meetings:
        mid = m['id']
        date_str = m['date']
        try:
            dt = datetime.strptime(date_str, '%Y-%m-%d')
            bulan_id = ['Januari','Februari','Maret','April','Mei','Juni',
                        'Juli','Agustus','September','Oktober','November','Desember']
            date_fmt = f"{dt.day} {bulan_id[dt.month-1]} {dt.year}"
        except:
            date_fmt = date_str

        add_para(doc, f'Pertemuan ({date_fmt})', size=12, align='center', space_before=6, space_after=3)

        # Build rows data
        rows_data = []
        for ak, al in zip(aspect_keys, aspect_labels):
            entries = assessments.get(mid, {}).get(ak, [])
            if not entries:
                entries = [{'kegiatan': '-', 'score': ''}]
            for i, entry in enumerate(entries):
                aspek_text = al if i == 0 else ''
                score = entry['score']
                rows_data.append({
                    'aspek': aspek_text,
                    'kegiatan': entry['kegiatan'],
                    'BB': '✓' if score == 'Belum Berkembang' else '',
                    'MB': '✓' if score == 'Mulai Berkembang' else '',
                    'BSH': '✓' if score == 'Berkembang Sesuai Harapan' else '',
                    'BSB': '✓' if score == 'Berkembang Sangat Baik' else '',
                })

        table = doc.add_table(rows=1 + len(rows_data), cols=6)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table)

        # Header row with blue shading
        headers = ['Aspek', 'Kegiatan', 'BB', 'MB', 'BSH', 'BSB']
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            set_cell_shading(cell, header_color)
            set_cell_width_twips(cell, col_widths[i])

        # Data rows
        for row_idx, rd in enumerate(rows_data):
            row = table.rows[row_idx + 1]
            
            # Aspek
            cell = row.cells[0]
            cell.text = ''
            p = cell.paragraphs[0]
            if rd['aspek']:
                run = p.add_run(rd['aspek'])
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            set_cell_width_twips(cell, col_widths[0])

            # Kegiatan
            cell = row.cells[1]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(rd['kegiatan'])
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            set_cell_width_twips(cell, col_widths[1])

            # BB, MB, BSH, BSB
            for col_idx, key in enumerate(['BB', 'MB', 'BSH', 'BSB'], 2):
                cell = row.cells[col_idx]
                cell.text = ''
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                val = rd[key]
                if val:
                    run = p.add_run(val)
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                set_cell_width_twips(cell, col_widths[col_idx])

    # === KETERANGAN STIMULASI ===
    add_para(doc, '', space_after=3)
    add_para(doc, 'Keterangan Stimulasi', size=12, align='center', space_before=6, space_after=3)

    keterangan = [
        ('Membaca', 'Kegiatan pengenalan huruf atau belajar teknis membaca (suku kata, kata, frasa, kalimat)'),
        ('Berhitung', 'Kegiatan pengenalan angka, konsep bilangan, konsep dasar matematika, dan operasi hitung matematika sederhana'),
        ('Menulis', 'Kegiatan menguatkan otot jari tangan (motoric halus) menggunakan alat tulis'),
        ('Brain Exercise', 'Kegiatan untuk menstimulasi kemampuan kognitif dan bahasa'),
        ('Sensory play', 'Kegiatan untuk menstimulasi koordinasi mata dengan tangan, panca indera, melatih fokus dan konsentrasi'),
        ('Kreativitas', 'Kegiatan untuk mengembangkan imajinasi dan keterampilan seni'),
    ]
    for label, desc in keterangan:
        add_para(doc, label, size=12, space_after=0)
        add_para(doc, desc, size=12, space_after=3)

    # === KETERANGAN SKALA ===
    add_para(doc, 'Keterangan Stimulasi', size=12, align='center', space_before=6, space_after=3)

    skala = [
        ('1', 'Belum Berkembang (BB)', 'Belum ingin berkegiatan'),
        ('2', 'Mulai Berkembang (MB)', 'Mulai ingin berkegiatan tetapi perlu stimulasi lebih lanjut'),
        ('3', 'Berkembang Sesuai Harapan (BSH)', 'Mampu berkegiatan dan terstimulasi dengan baik'),
        ('4', 'Berkembang Sangat Baik (BSB)', 'Mahir berkegiatan dan terstimulasi dengan sangat baik'),
    ]
    for num, label, desc in skala:
        add_para(doc, num, size=12, space_after=0)
        add_para(doc, label, size=12, space_after=0)
        add_para(doc, desc, size=12, space_after=3)

    # === KESIMPULAN STIMULASI ===
    add_para(doc, 'Kesimpulan Stimulasi', size=12, align='center', space_before=6, space_after=3)

    for idx, (ak, al) in enumerate(zip(aspect_keys, aspect_labels)):
        s = summary[ak]
        if s['Total'] == 0:
            continue
        levels = [
            ('Belum Berkembang', s['BB']),
            ('Mulai Berkembang', s['MB']),
            ('Berkembang Sesuai Harapan', s['BSH']),
            ('Berkembang Sangat Baik', s['BSB']),
        ]
        dominant = max(levels, key=lambda x: x[1])

        # Formulate rich narrative based on entries
        kegs = s['kegiatan_list']
        if len(kegs) > 1:
            keg_narrative = ", ".join(kegs[:-1]) + ", dan " + kegs[-1]
        elif len(kegs) == 1:
            keg_narrative = kegs[0]
        else:
            keg_narrative = "aktivitas stimulasi harian"

        status_text = dominant[0]
        if status_text == 'Belum Berkembang':
            prog_narrative = "belum ingin berkegiatan dan masih memerlukan bimbingan serta stimulasi yang lebih intensif."
        elif status_text == 'Mulai Berkembang':
            prog_narrative = "mulai menunjukkan ketertarikan untuk berkegiatan, namun masih memerlukan bantuan dan stimulasi lebih lanjut."
        elif status_text == 'Berkembang Sesuai Harapan':
            prog_narrative = "mampu mengikuti kegiatan dengan baik, aktif berpartisipasi, dan terstimulasi sesuai target perkembangan."
        else: # Berkembang Sangat Baik
            prog_narrative = "sangat mahir dalam berkegiatan, menunjukkan antusiasme yang tinggi, dan terstimulasi dengan sangat maksimal."

        final_text = f"{idx+1}. Dalam kegiatan {al}, {student['name']} telah belajar dan mempraktikkan: {keg_narrative}. Secara keseluruhan, {student['name']} {prog_narrative} Maka kemampuan {student['name']} dalam kegiatan {al} dinyatakan {status_text}."
        add_para(doc, final_text, size=11, space_after=4)

    # === TANDA TANGAN ===
    today = datetime.now()
    bulan_id = ['Januari','Februari','Maret','April','Mei','Juni',
                'Juli','Agustus','September','Oktober','November','Desember']
    date_str = f'{today.day} {bulan_id[today.month-1]} {today.year}'
    
    add_para(doc, '', space_after=3)
    add_para(doc, f'Nganjuk, {date_str}', size=12, space_after=6)

    sig_table = doc.add_table(rows=4, cols=2)
    remove_table_borders(sig_table)
    
    for i, label in enumerate(['Pengajar', 'Founder']):
        cell = sig_table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(label)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    sig_table.rows[1].height = Twips(600)
    sig_table.rows[2].height = Twips(600)

    # Get teacher name from args
    teacher_name = sys.argv[4] if len(sys.argv) > 4 else "Guru Pengajar"

    for i, name in enumerate([teacher_name, 'Admin']):
        cell = sig_table.rows[3].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(name)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    doc.save(output_path)
    print(output_path)

if __name__ == '__main__':
    main()
